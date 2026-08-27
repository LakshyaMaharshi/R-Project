"""Render dissertation_draft.md -> Dissertation_BIO-7057X_DRAFT.docx.

Formatting follows the BIO-7057X brief: Arial 11pt, 1.5 line spacing, left-aligned,
page numbers in the footer, named sections, Harvard reference list with hanging indent.
Re-run after editing the markdown. Not a general Markdown parser -- it understands the
specific structure of this draft.

Typesetting notes, since a few of these are easy to get wrong and hard to spot:

  * A4, not Letter. python-docx defaults to US Letter, which is the wrong paper for a
    UK submission and silently reflows every page.
  * The word count on the title page is INJECTED, not typed. Stating the count changes
    the count, so the document is built twice: once to measure, once with the measured
    numbers written in. It used to be maintained by hand and had gone stale.
  * Tables are ruled horizontally only, in the style journals use. Word's built-in grid
    styles put a coloured box round every cell, which reads as a spreadsheet.
  * Headings, figures and captions carry keep-with-next so nothing is orphaned at the
    foot of a page.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
SRC = HERE / "dissertation_draft.md"
OUT = HERE.parent / "Dissertation_BIO-7057X_DRAFT.docx"

# Sections whose headings must not be counted as main text. The brief excludes the
# reference list and the appendix; the contents page is our own scaffolding.
NON_BODY_HEADINGS = ("TABLE OF CONTENTS", "REFERENCES", "APPENDIX")
BODY_SECTIONS = ("1. INTRODUCTION", "2. MATERIALS", "3. RESULTS",
                 "4. DISCUSSION", "5. CONCLUSION")


# --------------------------------------------------------------------------- styles
def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(11)
    # ensure Arial applies to complex/east-asian script slots too
    rpr = st.element.get_or_add_rPr(); rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), "Arial")
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(10)
    pf.widow_control = True          # never leave a single line stranded across a page
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT   # the brief asks for left-aligned, not justified

    # Headings: black Arial, generous space above, and glued to the text beneath them.
    for name, size, before, after in (("Heading 1", 15, 20, 10),
                                      ("Heading 2", 12.5, 15, 7),
                                      ("Heading 3", 11.5, 12, 6)):
        s = doc.styles[name]
        s.font.name = "Arial"; s.font.size = Pt(size); s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        srpr = s.element.get_or_add_rPr(); srf = srpr.get_or_add_rFonts()
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            srf.set(qn(a), "Arial")
        hp = s.paragraph_format
        hp.space_before = Pt(before); hp.space_after = Pt(after)
        hp.line_spacing_rule = WD_LINE_SPACING.SINGLE
        hp.keep_with_next = True
        hp.keep_together = True


def setup_page(section):
    """A4 with even margins. python-docx defaults to Letter, which is simply wrong here."""
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.5)


# --------------------------------------------------------------------------- fields
def _field(run, instr):
    """Insert a Word field (PAGE, NUMPAGES, TOC ...) into a run."""
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, t, sep, e):
        run._r.append(el)


def add_footer_page_numbers(section):
    """'Page X of Y', centred. Linked footers would repeat the title page's blank one."""
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Page "); r.font.size = Pt(9); r.font.name = "Arial"
    _field(p.add_run(), "PAGE")
    r = p.add_run(" of "); r.font.size = Pt(9); r.font.name = "Arial"
    _field(p.add_run(), "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(9); run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table_of_contents(doc):
    h = doc.add_heading("Table of Contents", level=1)
    h.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    # \o "1-3" = heading levels 1-3, \h = hyperlinks, \z = hide tab leader in web view,
    # \u = use outline levels. Word fills this in on open (or with F9).
    _field(p.add_run(), 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    r = note.add_run("(Right-click and choose “Update Field” to refresh page numbers.)")
    r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    note.paragraph_format.space_after = Pt(0)


# --------------------------------------------------------------------------- tables
def _set_cell_bg(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table):
    """Horizontal rules only - the convention in journals, and far calmer than a grid."""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    spec = {"top": ("single", 12), "bottom": ("single", 12),
            "left": ("none", 0), "right": ("none", 0),
            "insideH": ("single", 4), "insideV": ("none", 0)}
    for edge, (val, sz) in spec.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val); el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "595959")
        borders.append(el)
    tblPr.append(borders)


def _repeat_header_row(row):
    """Header repeats when a table breaks across pages."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true")
    trPr.append(el)


NUMERIC_RE = re.compile(r"^[\s−–<>=~+\-0-9.,%()]+$")

def add_table(doc, rows):
    n_cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = "Table Grid"          # neutral base; the real look comes from the borders below
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _set_table_borders(t)

    # Columns that are entirely numbers read better centred than ragged-left.
    numeric_cols = set()
    for ci in range(n_cols):
        vals = [r[ci] for r in rows[1:] if ci < len(r) and r[ci].strip()]
        if vals and all(NUMERIC_RE.match(v) for v in vals):
            numeric_cols.add(ci)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci >= n_cols:
                continue
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_before = Pt(3); pf.space_after = Pt(3)
            if ri > 0 and ci in numeric_cols:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
            if ri == 0:
                run.bold = True
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in numeric_cols \
                               else WD_ALIGN_PARAGRAPH.LEFT
                _set_cell_bg(cell, "F2F2F2")
    _repeat_header_row(t.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 1.0


# --------------------------------------------------------------------------- inline
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITAL_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")

def add_runs(p, text):
    """Add text to a paragraph, honouring **bold** and *italic* spans."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|(?<!\*)\*([^*]+)\*(?!\*)", text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            p.add_run(m.group(1)).bold = True
        else:
            p.add_run(m.group(2)).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def strip_comments(md):
    return re.sub(r"<!--.*?-->", "", md, flags=re.DOTALL)

SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")
IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
CAP_RE = re.compile(r"\*(Table|Figure)\s")


def add_image(doc, rel_path, width_in=6.1):
    """Insert a figure, centred, glued to its caption."""
    path = (SRC.parent / rel_path).resolve()
    if not path.exists():
        raise SystemExit(f"figure not found: {path}\n  (run: Rscript code/09_figures.R)")
    doc.add_picture(str(path), width=Inches(width_in))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(10); pf.space_after = Pt(4)
    pf.keep_with_next = True        # the caption must not land on the next page alone
    pf.line_spacing = 1.0


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(2); pf.space_after = Pt(14)
    run = p.add_run(text.strip("*"))
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")

def parse_row(line):
    cells = line.strip().strip("|").split("|")
    return [BOLD_RE.sub(r"\1", c.strip()) for c in cells]


# --------------------------------------------------------------------------- title
def add_title_page(doc, block, counts):
    """Title page. The word count is injected from `counts`, never typed by hand."""
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    for j, b in enumerate(block):
        txt = BOLD_RE.sub(r"\1", b)

        # Substitute the measured counts into whichever lines quote them.
        if counts:
            txt = re.sub(r"(Word count \(main text\):\s*)[\d,]+",
                         lambda m: m.group(1) + f"{counts['main']:,}", txt)
            txt = re.sub(r"(whole-document count, which includes all of those, is\s*)[\d,]+",
                         lambda m: m.group(1) + f"{counts['total']:,}", txt)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.3

        if j == 0:                                   # the title
            r = p.add_run(txt); r.bold = True; r.font.size = Pt(17)
            pf.space_after = Pt(30); pf.line_spacing = 1.25
        elif txt.startswith("("):                    # the word-count explanatory note
            r = p.add_run(txt); r.font.size = Pt(9); r.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            pf.space_before = Pt(16); pf.line_spacing = 1.15
        elif j == 1:                                 # author
            r = p.add_run(txt); r.font.size = Pt(13); r.bold = True
            pf.space_after = Pt(4)
        else:
            r = p.add_run(txt); r.font.size = Pt(11)
            pf.space_after = Pt(3)


# --------------------------------------------------------------------------- build
def build(counts=None):
    """Render the document. Pass `counts` on the second pass to fill in the title page."""
    md = strip_comments(SRC.read_text(encoding="utf-8"))
    doc = Document()
    set_base_style(doc)
    setup_page(doc.sections[0])

    lines = md.splitlines()
    i = 0
    body_started = False
    while i < len(lines):
        raw = lines[i]; line = raw.strip()
        if line == "" or line == "---":
            i += 1; continue

        if line.startswith("# "):
            name = line[2:].strip()
            up = name.upper()

            if up == "TITLE PAGE":
                i += 1
                block = []
                while i < len(lines) and not lines[i].strip().startswith("# "):
                    if lines[i].strip() not in ("", "---"):
                        block.append(lines[i].strip())
                    i += 1
                add_title_page(doc, block, counts)

                # A new section, so the title page can stay free of a page number while
                # everything after it is numbered.
                body = doc.add_section(WD_SECTION.NEW_PAGE)
                setup_page(body)
                add_footer_page_numbers(body)
                add_table_of_contents(doc)
                doc.add_page_break()
                body_started = True
                continue

            if up.startswith("REFERENCES") or up.startswith("APPENDIX"):
                doc.add_page_break()
            doc.add_heading(name, level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith(">"):
            i += 1; continue     # editorial note, skip in docx

        m_img = IMG_RE.match(line)
        if m_img:
            add_image(doc, m_img.group(2))
            i += 1; continue

        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [parse_row(r) for r in table_lines if not SEP_RE.match(r)]
            if rows:
                add_table(doc, rows)  # table contents excluded from word count per brief
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
            p.paragraph_format.space_after = Pt(4)
        elif CAP_RE.match(line):
            add_caption(doc, line)
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    apply_reference_hanging_indent(doc)
    return doc


def apply_reference_hanging_indent(doc):
    """Harvard reference list: hanging indent, single-spaced, as the brief asks.

    This was described in a comment for a long time but never actually implemented, so
    the reference list rendered as ordinary body paragraphs.
    """
    in_refs = False
    n = 0
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            up = p.text.strip().upper()
            in_refs = up.startswith("REFERENCES")
            continue
        if not in_refs or not p.text.strip():
            continue
        pf = p.paragraph_format
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        pf.line_spacing = 1.15
        pf.space_after = Pt(8)
        n += 1
    return n


def save(doc):
    out = OUT
    try:
        doc.save(out)
    except PermissionError:
        out = OUT.with_name(OUT.stem + "_UPDATED" + OUT.suffix)
        doc.save(out)
        print(f"WARNING: '{OUT.name}' is locked (open in Word?). Saved to '{out.name}' instead.")
        print("        Close Word and re-run to update the original filename.")
    return out


def main():
    # Two passes. Stating the word count on the title page changes the document, so the
    # count is measured on a first build and written into the second.
    tmp = HERE / "_wordcount_probe.docx"   # deleted below; never committed
    build().save(str(tmp))
    counts = measure(Document(str(tmp)))
    tmp.unlink(missing_ok=True)

    doc = build(counts)
    out = save(doc)
    print(f"WROTE {out}")
    report(measure(Document(str(out))))


# --------------------------------------------------------------------------- counting
def measure(doc):
    """Count the way the assessment brief defines it.

    Word counts the whole document. The brief excludes table contents, table titles,
    figure legends and the reference list, so Word's number is always noticeably higher.
    """
    def wc(s):
        return len([t for t in re.split(r"\s+", s.strip()) if t])

    tables = sum(wc(c.text) for t in doc.tables for row in t.rows for c in row.cells)
    title_pg = captions = refs = appendix = body = headings = front = 0
    in_refs = in_appx = in_front = False
    started = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        if p.style.name.startswith("Heading"):
            up = txt.upper()
            in_refs = up.startswith("REFERENCES")
            in_appx = up.startswith("APPENDIX")
            in_front = up.startswith("TABLE OF CONTENTS")
            started = True
            if in_refs:
                refs += wc(txt)
            elif in_appx:
                appendix += wc(txt)
            elif in_front:
                front += wc(txt)
            else:
                headings += wc(txt)
            continue
        if not started:
            title_pg += wc(txt); continue
        if in_front:
            front += wc(txt); continue
        if re.match(r"^(Table|Figure)\s+\d", txt) or re.match(r"^Table\s+[A-E]\d", txt):
            captions += wc(txt); continue
        if in_appx:
            appendix += wc(txt); continue
        if in_refs:
            refs += wc(txt); continue
        body += wc(txt)

    main_text = body + headings
    return dict(body=body, headings=headings, title=title_pg, tables=tables,
                captions=captions, refs=refs, appendix=appendix, front=front,
                main=main_text,
                total=main_text + title_pg + tables + captions + refs + appendix + front)


def report(c):
    print(f"  body text              {c['body']:>6}")
    print(f"  headings               {c['headings']:>6}")
    print(f"  title page             {c['title']:>6}  (excluded)")
    print(f"  contents page          {c['front']:>6}  (excluded)")
    print(f"  table contents         {c['tables']:>6}  (excluded by the brief)")
    print(f"  table/figure captions  {c['captions']:>6}  (excluded by the brief)")
    print(f"  reference list         {c['refs']:>6}  (excluded by the brief)")
    print(f"  appendix               {c['appendix']:>6}  (excluded by the brief)")
    print(f"  MAIN TEXT              {c['main']:>6}  <- printed on the title page")
    print(f"  whole document         {c['total']:>6}  <- what Word will display")
    if c["main"] > 8000:
        print("  ** OVER 8,000 - the brief requires the Module Organiser's permission **")
    elif c["main"] < 6000:
        print("  ** UNDER 6,000 - below the brief's minimum **")
    else:
        print("  within the brief's 6,000-8,000 range")


if __name__ == "__main__":
    main()
